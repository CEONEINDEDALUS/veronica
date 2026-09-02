"""
Extracts plain text from many document types so they can be chunked & embedded.
"""
from __future__ import annotations
import os
import csv
from pathlib import Path

SUPPORTED_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".log",
    ".pdf", ".docx", ".csv", ".tsv", ".xlsx", ".xls",
    ".html", ".htm", ".json", ".py", ".js", ".ts", ".java", ".c", ".cpp", ".rs", ".go",
}

MAX_FILE_BYTES = 64 * 1024 * 1024
MAX_EXTRACTED_CHARS = 4_000_000
MAX_SHEETS = 20
MAX_CSV_ROWS = 200_000


class LoadError(Exception):
    pass


def load_file(path: str) -> str:
    ext = Path(path).suffix.lower()
    try:
        size = os.path.getsize(path)
        if size > MAX_FILE_BYTES:
            raise LoadError(f"File too large ({size / (1024 * 1024):.0f} MB, limit is "
                            f"{MAX_FILE_BYTES // (1024 * 1024)} MB)")
        if size == 0:
            raise LoadError("Empty file")
        if ext in (".txt", ".md", ".markdown", ".log", ".json", ".py", ".js", ".ts",
                   ".java", ".c", ".cpp", ".rs", ".go"):
            return _load_plain_text(path)
        if ext == ".pdf":
            return _load_pdf(path)
        if ext == ".docx":
            return _load_docx(path)
        if ext in (".csv", ".tsv"):
            return _load_csv(path, delimiter="," if ext == ".csv" else "\t")
        if ext in (".xlsx", ".xls"):
            return _load_excel(path)
        if ext in (".html", ".htm"):
            return _load_html(path)
        # Unknown extension: try as text, best effort
        return _load_plain_text(path)
    except LoadError:
        raise
    except Exception as e:
        raise LoadError(f"Failed to read {os.path.basename(path)}: {e}")


def _cap(text: str) -> str:
    if len(text) > MAX_EXTRACTED_CHARS:
        return text[:MAX_EXTRACTED_CHARS] + "\n\n[Document truncated: exceeded extraction limit]"
    return text


def _load_plain_text(path: str) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return _cap(f.read())
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise LoadError("Could not decode text file with common encodings")


def _load_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    parts = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            parts.append(f"[Page {i + 1}]\n{text}")
    return _cap("\n\n".join(parts))


def _load_docx(path: str) -> str:
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return _cap("\n".join(parts))


def _load_csv(path: str, delimiter: str) -> str:
    rows = []
    with open(path, "r", encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.reader(f, delimiter=delimiter)
        header = next(reader, None)
        if header:
            rows.append(" | ".join(header))
        for row in reader:
            rows.append(" | ".join(row))
            if len(rows) >= MAX_CSV_ROWS:
                rows.append("[Truncated: row limit reached]")
                break
    return _cap("\n".join(rows))


def _load_excel(path: str) -> str:
    import pandas as pd
    xls = pd.ExcelFile(path)
    parts = []
    for sheet in xls.sheet_names[:MAX_SHEETS]:
        df = xls.parse(sheet)
        parts.append(f"[Sheet: {sheet}]")
        parts.append(df.to_csv(index=False))
    return _cap("\n".join(parts))


def _load_html(path: str) -> str:
    from bs4 import BeautifulSoup
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    return _cap(soup.get_text(separator="\n"))


def iter_supported_files(root: str):
    for dirpath, _dirnames, filenames in os.walk(root):
        for name in filenames:
            if Path(name).suffix.lower() in SUPPORTED_EXTENSIONS:
                yield os.path.join(dirpath, name)
